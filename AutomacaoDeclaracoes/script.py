import win32print
import win32api
from docx import Document
import os
from datetime import datetime
import locale
from time import sleep

with open("nomes.txt", "r", encoding="utf-8") as arquivo:
    names = [linha.strip() for linha in arquivo if linha.strip()]

os.makedirs("Declaracoes", exist_ok=True)

cidade = input("Digite a cidade da instituição de ensino: ").title()

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except:
    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')

hoje = datetime.today()
placeholders = {
    "[CIDADE]": cidade,
    "[DIA]": str(hoje.day),
    "[MES]": hoje.strftime('%B').title(),
    "[ANO]": str(hoje.year)
}

for name in names:
    doc = Document("Declaração Escolar.docx")

    for paragrafo in doc.paragraphs:
        texto = paragrafo.text

        for placeholder, valor in placeholders.items():
            if placeholder in texto:
                texto = texto.replace(placeholder, valor)
                paragrafo.clear()
                paragrafo.add_run(texto)

        if "[NOME]" in texto:
            partes = texto.split("[NOME]")
            paragrafo.clear()
            paragrafo.add_run(partes[0])

            run_nome = paragrafo.add_run(name)
            run_nome.bold = True
            run_nome.underline = True

            if len(partes) > 1:
                paragrafo.add_run(partes[1])

    nome_arquivo = f"Declaracoes/Declaracao_{name.replace(' ', '')}.docx"
    doc.save(nome_arquivo)
    print(f"Salvo: {nome_arquivo}")

printers = win32print.EnumPrinters(2)
if printers:
    print("Impressoras disponíveis:")
    for i, p in enumerate(printers):
        print(f"[{i}] {p[2]}")

    indice = int(input("Digite o número da impressora desejada: "))
    printer_name = printers[indice][2]

    win32print.SetDefaultPrinter(printer_name)
    print(f"Impressora definida: {printer_name}")

    pasta = "Declaracoes"
    arquivos = [declaracao for declaracao in os.listdir(pasta) if declaracao.endswith(".docx")]

    for arquivo in arquivos:
        caminho = os.path.abspath(os.path.join(pasta, arquivo))
        print(f"Imprimindo: {caminho}")
        win32api.ShellExecute(0, "print", caminho, None, pasta, 0)
        sleep(20) 
else:
    print("Nenhuma impressora detectada.")
    